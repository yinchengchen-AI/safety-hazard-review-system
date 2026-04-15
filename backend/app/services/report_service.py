import io
import os
import tempfile
import uuid
import base64
from datetime import datetime, timezone
from zoneinfo import ZoneInfo
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.async_api import async_playwright
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from PIL import Image as PILImage
from app.models import ReviewTask, TaskHazard, Photo, Report, Hazard
from app.services.storage_service import StorageService

MAX_REPORT_PHOTO_SIZE = 5 * 1024 * 1024  # 5MB

STATUS_MAP = {
    "pending": "待复核",
    "passed": "已通过",
    "failed": "未通过",
}


class ReportService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.storage = StorageService()

    async def generate_report(self, task_id: uuid.UUID) -> Report:
        result = await self.db.execute(
            select(ReviewTask)
            .where(ReviewTask.id == task_id)
            .options(selectinload(ReviewTask.creator))
        )
        task = result.scalar_one_or_none()
        if not task:
            raise ValueError("Review task not found")

        # Update report status to processing
        report_result = await self.db.execute(select(Report).where(Report.task_id == task_id))
        report = report_result.scalar_one_or_none()
        if not report:
            report = Report(task_id=task_id, status="pending")
            self.db.add(report)
            await self.db.flush()

        report.status = "processing"
        await self.db.commit()

        try:
            # Fetch task hazards with photos
            th_result = await self.db.execute(
                select(TaskHazard)
                .where(TaskHazard.task_id == task_id)
                .options(
                    selectinload(TaskHazard.hazard).selectinload(Hazard.enterprise),
                    selectinload(TaskHazard.photos),
                    selectinload(TaskHazard.reviewer),
                )
            )
            task_hazards = th_result.scalars().all()

            # Load photo binaries
            for th in task_hazards:
                th._photo_data = []
                for photo in th.photos:
                    if photo.deleted_at is not None:
                        continue
                    try:
                        data = self.storage.get_file(photo.original_path)
                        content = data.read()
                        data.close()
                        data.release_conn()
                        th._photo_data.append((photo.mime_type or "image/jpeg", content))
                    except Exception:
                        pass

            # Generate Word
            word_path = await self._generate_word(task, task_hazards)

            # Generate PDF via Playwright
            pdf_path = await self._generate_pdf(task, task_hazards)

            report.word_path = word_path
            report.pdf_path = pdf_path
            report.status = "completed"
            report.generated_at = datetime.now(ZoneInfo("Asia/Shanghai"))
        except Exception as e:
            report.status = "failed"
            report.error_message = str(e)

        await self.db.commit()
        await self.db.refresh(report)
        return report

    def _compress_photo_if_needed(self, content: bytes, mime_type: str) -> bytes:
        if len(content) <= MAX_REPORT_PHOTO_SIZE:
            return content
        try:
            img = PILImage.open(io.BytesIO(content))
            # Resize if very large
            max_dimension = 1920
            if max(img.size) > max_dimension:
                ratio = max_dimension / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, PILImage.Resampling.LANCZOS)
            # Convert to RGB for JPEG output
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            buffer = io.BytesIO()
            quality = 95
            while quality >= 60:
                buffer.seek(0)
                buffer.truncate()
                img.save(buffer, format="JPEG", quality=quality, optimize=True)
                if buffer.tell() <= MAX_REPORT_PHOTO_SIZE:
                    break
                quality -= 5
            return buffer.getvalue()
        except Exception:
            return content

    async def _generate_word(self, task: ReviewTask, task_hazards: list) -> str:
        doc = Document()

        # Title
        title = doc.add_heading("安全生产隐患复核报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Task info
        doc.add_paragraph(f"任务名称: {task.name}")
        doc.add_paragraph(f"创建人: {task.creator.username if task.creator else ''}")
        doc.add_paragraph(f"创建时间: {task.created_at.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        for idx, th in enumerate(task_hazards, 1):
            hazard = th.hazard
            enterprise = hazard.enterprise if hazard else None

            # Section heading
            doc.add_heading(f"隐患 {idx}", level=2)

            # Basic info table
            info_table = doc.add_table(rows=0, cols=2)
            info_table.style = "Light Grid Accent 1"

            def add_info_row(label, value):
                row = info_table.add_row().cells
                row[0].text = label
                row[1].text = value or "-"

            reporting_unit = None
            if hazard:
                reporting_unit = hazard.reporting_unit
                if not reporting_unit and hazard.batch:
                    reporting_unit = hazard.batch.reporting_unit

            add_info_row("隐患描述", hazard.content if hazard else None)
            add_info_row("位置", hazard.location if hazard else None)
            add_info_row("企业名称", enterprise.name if enterprise else None)
            add_info_row("所属地区", enterprise.region if enterprise else None)
            add_info_row("地址", enterprise.address if enterprise else None)
            add_info_row("联系人", enterprise.contact_person if enterprise else None)
            add_info_row("行业领域", enterprise.industry_sector if enterprise else None)
            add_info_row("企业类型", enterprise.enterprise_type if enterprise else None)
            add_info_row("上报单位", reporting_unit)
            add_info_row("复核结论", th.conclusion)
            add_info_row("复核状态", STATUS_MAP.get(th.status_in_task, th.status_in_task or "待复核"))
            add_info_row("复核人", th.reviewer.username if th.reviewer else None)
            add_info_row("复核时间", th.reviewed_at.strftime("%Y-%m-%d %H:%M") if th.reviewed_at else None)

            # Photos
            photo_data = getattr(th, "_photo_data", [])
            if photo_data:
                doc.add_paragraph("复核照片:")
                for mime_type, content in photo_data:
                    compressed = self._compress_photo_if_needed(content, mime_type)
                    ext = "jpg"  # compressed photos are JPEG
                    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
                        tmp.write(compressed)
                        tmp_path = tmp.name
                    try:
                        doc.add_picture(tmp_path, width=Inches(3.0))
                    finally:
                        os.unlink(tmp_path)

            doc.add_paragraph()

        # Save to temp file and upload
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        with open(tmp_path, "rb") as f:
            data = f.read()

        object_name = f"reports/{task.id}/{uuid.uuid4()}.docx"
        self.storage.client.put_object(
            self.storage.bucket,
            object_name,
            io.BytesIO(data),
            length=len(data),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        os.unlink(tmp_path)
        return f"/{self.storage.bucket}/{object_name}"

    async def _generate_pdf(self, task: ReviewTask, task_hazards: list) -> str:
        html = self._build_html(task, task_hazards)

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.set_content(html)
            await page.pdf(path=tmp_path, format="A4", margin={"top": "1cm", "bottom": "1cm", "left": "1cm", "right": "1cm"})
            await browser.close()

        with open(tmp_path, "rb") as f:
            data = f.read()

        object_name = f"reports/{task.id}/{uuid.uuid4()}.pdf"
        self.storage.client.put_object(
            self.storage.bucket,
            object_name,
            io.BytesIO(data),
            length=len(data),
            content_type="application/pdf",
        )
        os.unlink(tmp_path)
        return f"/{self.storage.bucket}/{object_name}"

    def _build_html(self, task: ReviewTask, task_hazards: list) -> str:
        rows = ""
        for idx, th in enumerate(task_hazards, 1):
            hazard = th.hazard
            enterprise = hazard.enterprise if hazard else None

            photos_html = ""
            for mime_type, content in getattr(th, "_photo_data", []):
                compressed = self._compress_photo_if_needed(content, mime_type)
                b64 = base64.b64encode(compressed).decode("utf-8")
                data_uri = f"data:image/jpeg;base64,{b64}"
                photos_html += f'<img src="{data_uri}" style="max-width:200px;max-height:200px;margin:4px;border:1px solid #ccc;" />'

            if not photos_html:
                photos_html = "-"

            reviewer = th.reviewer.username if th.reviewer else "-"
            reviewed_at = th.reviewed_at.strftime("%Y-%m-%d %H:%M") if th.reviewed_at else "-"
            status_text = STATUS_MAP.get(th.status_in_task, th.status_in_task or "待复核")

            reporting_unit = ""
            if hazard:
                reporting_unit = hazard.reporting_unit or ""
                if not reporting_unit and hazard.batch:
                    reporting_unit = hazard.batch.reporting_unit or ""

            rows += f"""
            <tr>
                <td>{idx}</td>
                <td>{hazard.content or ""}</td>
                <td>{hazard.location or ""}</td>
                <td>{enterprise.name if enterprise else ""}</td>
                <td>{enterprise.region if enterprise else ""}</td>
                <td>{reporting_unit}</td>
                <td>{th.conclusion or ""}</td>
                <td>{status_text}</td>
                <td>{reviewer}</td>
                <td>{reviewed_at}</td>
                <td>{photos_html}</td>
            </tr>
            """

        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: "Noto Sans CJK SC", "Microsoft YaHei", sans-serif; font-size: 12px; }}
                h1 {{ text-align: center; }}
                h2 {{ margin-top: 24px; }}
                table {{ width: 100%; border-collapse: collapse; margin-top: 12px; }}
                th, td {{ border: 1px solid #333; padding: 8px; text-align: left; vertical-align: top; }}
                th {{ background: #f2f2f2; }}
                .info {{ margin: 8px 0; }}
                .info strong {{ display: inline-block; width: 120px; }}
            </style>
        </head>
        <body>
            <h1>安全生产隐患复核报告</h1>
            <div class="info"><strong>任务名称:</strong> {task.name}</div>
            <div class="info"><strong>创建人:</strong> {task.creator.username if task.creator else ""}</div>
            <div class="info"><strong>创建时间:</strong> {task.created_at.strftime('%Y-%m-%d %H:%M')}</div>

            <table>
                <tr>
                    <th>序号</th>
                    <th>隐患描述</th>
                    <th>位置</th>
                    <th>企业名称</th>
                    <th>所属地区</th>
                    <th>上报单位</th>
                    <th>复核结论</th>
                    <th>复核状态</th>
                    <th>复核人</th>
                    <th>复核时间</th>
                    <th>复核照片</th>
                </tr>
                {rows}
            </table>
        </body>
        </html>
        """
